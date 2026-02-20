# ============================================================================
# CV EM WORD PARA VAGA JAVA PLENO - VERSÃO FIEL AO FORMATO PDF
# Formato idêntico ao CV ATS em PDF, apenas convertido para Word
# ============================================================================

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def criar_cv_java_word_fiel():
    """
    Cria um currículo em Word com exatamente o mesmo formato do CV ATS em PDF
    """
    doc = Document()
    
    # Configurar margens similares ao A4
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # CABEÇALHO (EXATAMENTE como no PDF)
    # Nome
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Benevanio Santos\n')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    # Título (idêntico ao do PDF)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Desenvolvedor Back-End | Java, Spring Boot & Arquitetura Distribuída\n')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    
    # Espaço
    doc.add_paragraph()
    
    # INFORMAÇÕES DE CONTATO (com dados do RH)
    p = doc.add_paragraph()
    
    # Linha 1: Endereço e contato (endereço atualizado)
    run = p.add_run('Porto da Folha, Sergipe, Brasil')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('benevaniosantos930@gmail.com')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('WhatsApp: +55 19 99828-3835')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    # Nova linha
    p = doc.add_paragraph()
    
    # Linha 2: LinkedIn, GitHub, Portfolio
    run = p.add_run('LinkedIn: linkedin.com/in/bene-tesla')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('GitHub: github.com/Benevanio')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('Portfólio: bene-tesla-dev.vercel.app')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    # Linha 3: Informações específicas do RH
    p = doc.add_paragraph()
    
    run = p.add_run('Data de Nascimento: 25/06/1998')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('Pretensão Salarial: R$ 6.500,00')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.bold = True  # Destaque na pretensão salarial
    
    p.add_run(' | ')
    
    run = p.add_run('Inglês: Iniciante (em estudo, leitura técnica)')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    # Espaçamento duplo (como no PDF)
    for _ in range(2):
        doc.add_paragraph()
    
    # ============================================
    # RESUMO PROFISSIONAL (IDÊNTICO AO PDF)
    # ============================================
    p = doc.add_paragraph()
    run = p.add_run('RESUMO PROFISSIONAL\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('''Desenvolvedor Back-End com sólida experiência em Java e ecossistema Spring, atuando desde versões corporativas (Java 6) até as mais modernas (Java 17+). Especializado no desenvolvimento de APIs RESTful, microsserviços e sistemas distribuídos com Spring Boot. Foco em qualidade de código, performance (otimização JVM), testes e boas práticas de arquitetura (Clean Architecture, DDD). Experiência prática em projetos de integração e alta disponibilidade.''')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    # Espaçamento
    for _ in range(2):
        doc.add_paragraph()
    
    # ============================================
    # EXPERIÊNCIA PROFISSIONAL (IDÊNTICO AO PDF)
    # ============================================
    p = doc.add_paragraph()
    run = p.add_run('EXPERIÊNCIA PROFISSIONAL\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    # Cargo e período
    p = doc.add_paragraph()
    run = p.add_run('SysMap Solutions – Desenvolvedor Back-End (Terceirizado Natura)\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('Jun 2023 – Atual | Remoto')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    # Itens da experiência (bullet points)
    experience_items = [
        "Desenvolvimento e manutenção de APIs RESTful e microsserviços utilizando Spring Boot, Spring MVC e Spring Data JPA, garantindo 99.9% de disponibilidade.",
        "Otimização de performance de aplicações Java através de análise de GC, tuning de JVM e melhorias em consultas ao banco de dados, resultando em um aumento de 25% no throughput.",
        "Implementação de logging estruturado e estratégias de monitoramento para sistemas distribuídos, aumentando a visibilidade operacional em 40%.",
        "Integração com bancos de dados relacionais (MySQL, PostgreSQL, Oracle) utilizando JPA/Hibernate, garantindo eficiência e consistência, com redução de 30% no tempo de consulta.",
        "Containerização de aplicações Java com Docker e deploy em ambientes cloud (AWS EC2), reduzindo o tempo de deploy em 40%.",
        "Configuração de pipelines CI/CD com Jenkins para build, teste e deploy automatizados, reduzindo o tempo de integração em 50%.",
        "Aplicação de princípios de design (SOLID) e arquiteturas limpas (Clean Architecture) para criação de sistemas sustentáveis, facilitando a manutenção e reduzindo bugs em 20%."
    ]
    
    for item in experience_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    # Espaçamento
    for _ in range(2):
        doc.add_paragraph()
    
    # ============================================
    # HABILIDADES TÉCNICAS (IDÊNTICO AO PDF)
    # ============================================
    p = doc.add_paragraph()
    run = p.add_run('HABILIDADES TÉCNICAS\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    # Formatar habilidades exatamente como no PDF
    skills_content = [
        ("Linguagens & Plataforma:", "Java (6, 8, 11, 17+), JVM"),
        ("Frameworks & Bibliotecas:", "Spring Boot, Spring MVC, Spring Data JPA, Hibernate, Maven"),
        ("APIs & Protocolos:", "RESTful APIs, JSON, XML, SOAP (conceitos)"),
        ("Bancos de Dados:", "MySQL, PostgreSQL, Oracle, H2, Redis (cache)"),
        ("Testes:", "JUnit, Mockito, Test Containers"),
        ("Ferramentas & DevOps:", "Docker, Kubernetes (conceitos), Terraform (conceitos), Jenkins, AWS EC2, Git, Postman, Kibana"),
        ("Conceitos & Arquitetura:", "Microsserviços, Clean Architecture, SOLID, DDD, Caching"),
        ("Soft Skills:", "Comunicação, Trabalho em Equipe, Resolução de Problemas")
    ]
    
    for category, skills in skills_content:
        p = doc.add_paragraph()
        run = p.add_run(f'{category} ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        
        run = p.add_run(skills)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    # Espaçamento
    for _ in range(2):
        doc.add_paragraph()
    
    # ============================================
    # PROJETOS RELEVANTES (IDÊNTICO AO PDF)
    # ============================================
    p = doc.add_paragraph()
    run = p.add_run('PROJETOS RELEVANTES\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    projetos = [
        ("RestSpring:", "API RESTful completa desenvolvida com Spring Boot, incluindo JPA/Hibernate, H2 para testes, documentação Swagger/OpenAPI e versionamento de API."),
        ("Sistema de Integração:", "Contribuição para um sistema de integração corporativa utilizando Java e Spring Boot para orquestração de fluxos de dados entre sistemas.")
    ]
    
    for projeto_nome, projeto_desc in projetos:
        p = doc.add_paragraph()
        run = p.add_run(f'{projeto_nome} ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        
        run = p.add_run(projeto_desc)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    # Espaçamento
    for _ in range(2):
        doc.add_paragraph()
    
    # ============================================
    # FORMAÇÃO ACADÊMICA (IDÊNTICO AO PDF)
    # ============================================
    p = doc.add_paragraph()
    run = p.add_run('FORMAÇÃO ACADÊMICA\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('PUC Minas')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    run = p.add_run(' – Pós-Graduação em Arquitetura de Software Distribuído (2025 – 2027, em andamento)')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('Anhanguera Educacional')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    run = p.add_run(' – Bacharelado em Engenharia de Software (2021 – 2025)')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    # Salvar documento
    data_atual = datetime.datetime.now().strftime('%Y%m%d')
    filename = f"CV_Java_Pleno_Benevanio_Santos_{data_atual}.docx"
    
    doc.save(filename)
    
    print("="*70)
    print("✅ CV EM WORD GERADO COM SUCESSO!")
    print("="*70)
    print(f"📄 Arquivo: {filename}")
    print("\n🎯 ESTRUTURA (IDÊNTICA AO PDF ORIGINAL):")
    print("   1. Cabeçalho com dados pessoais")
    print("   2. Resumo Profissional")
    print("   3. Experiência Profissional")
    print("   4. Habilidades Técnicas")
    print("   5. Projetos Relevantes")
    print("   6. Formação Acadêmica")
    print("\n📌 DADOS DO RH INCLUÍDOS NO CABEÇALHO:")
    print("   ✓ Endereço: Porto da Folha, Sergipe")
    print("   ✓ Data de Nascimento: 25/06/1998")
    print("   ✓ Pretensão Salarial: R$ 6.500,00")
    print("="*70)
    
    return filename

# ============================================================================
# VERSÃO MAIS ATUALIZADA (com experiência real em Java 6/7/8)
# ============================================================================

def criar_cv_java_atualizado():
    """
    Versão atualizada com sua experiência real em Java 6/7/8
    """
    doc = Document()
    
    # Configurar margens
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # CABEÇALHO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Benevanio Santos\n')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Desenvolvedor Java Pleno | 3 anos experiência | Java 6/7/8 | Spring Framework\n')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    
    doc.add_paragraph()
    
    # INFORMAÇÕES (com dados do RH)
    p = doc.add_paragraph()
    run = p.add_run('Porto da Folha, Sergipe, Brasil')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('benevaniosantos930@gmail.com')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('WhatsApp: +55 19 99828-3835')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    
    run = p.add_run('LinkedIn: linkedin.com/in/bene-tesla')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('GitHub: github.com/Benevanio')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    
    p.add_run(' | ')
    
    run = p.add_run('Data Nasc: 25/06/1998 | Pretensão: R$ 6.500,00')
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.bold = True
    
    for _ in range(2):
        doc.add_paragraph()
    
    # RESUMO ATUALIZADO
    p = doc.add_paragraph()
    run = p.add_run('RESUMO PROFISSIONAL\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('''Desenvolvedor Java Pleno com 3 anos de experiência prática em desenvolvimento e manutenção de sistemas corporativos. 
Experiência sólida com Java 6, 7, 8 e Spring Framework 3.x.x em ambiente de produção. 
Atuação em migração de sistemas legados para tecnologias modernas, com conhecimento em Oracle Database para cadastro de produtos e dados transacionais. 
Foco em qualidade de código, otimização de performance e boas práticas de desenvolvimento. 
Em processo de atualização para Java 17 e Spring Boot através de estudos contínuos.''')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    for _ in range(2):
        doc.add_paragraph()
    
    # EXPERIÊNCIA ATUALIZADA
    p = doc.add_paragraph()
    run = p.add_run('EXPERIÊNCIA PROFISSIONAL\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('SysMap Solutions – Desenvolvedor Back-End Java\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    p = doc.add_paragraph()
    run = p.add_run('Out 2021 – Jun 2024 (2 anos e 8 meses) | Remoto')
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    experiencia_items = [
        "Desenvolvimento e manutenção de sistemas corporativos com Java 6, 7, 8 e Spring Framework 3.x.x",
        "Otimização de performance em sistemas legados: redução de 40% no tempo de resposta",
        "Trabalho com Oracle Database para cadastro de produtos, modelagem de dados e otimização de queries",
        "Desenvolvimento de Web Services SOAP com JAX-WS para integração entre sistemas",
        "Containerização de aplicações legadas com Docker",
        "Suporte a sistemas em produção e correção de bugs críticos"
    ]
    
    for item in experiencia_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    for _ in range(2):
        doc.add_paragraph()
    
    # HABILIDADES ATUALIZADAS
    p = doc.add_paragraph()
    run = p.add_run('HABILIDADES TÉCNICAS\n')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    
    habilidades = [
        ("Java:", "6, 7, 8 (3 anos experiência corporativa) | Java 17 (em estudo)"),
        ("Spring:", "Framework 3.x.x (Spring MVC, Data, Security) | Spring Boot (transição)"),
        ("Banco de Dados:", "Oracle (experiência prática) | MySQL | PostgreSQL"),
        ("Web Services:", "SOAP (JAX-WS) | REST (conceitos)"),
        ("Ferramentas:", "Maven | Git | Jenkins | Docker | JUnit"),
        ("Conceitos:", "Clean Code | SOLID | Design Patterns | Performance Tuning")
    ]
    
    for categoria, desc in habilidades:
        p = doc.add_paragraph()
        run = p.add_run(f'{categoria} ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    # Salvar
    data_atual = datetime.datetime.now().strftime('%Y%m%d')
    filename = f"CV_Java_Atualizado_Benevanio_{data_atual}.docx"
    doc.save(filename)
    
    print(f"\n📋 CV ATUALIZADO GERADO: {filename}")
    print("   (Com experiência real em Java 6/7/8 e Spring 3.x.x)")
    
    return filename

# ============================================================================
# EXECUÇÃO PRINCIPAL
# ============================================================================

if __name__ == "__main__":
    print("="*70)
    print("GERADOR DE CV JAVA PLENO - MICROSOFT WORD")
    print("="*70)
    
    print("\nEscolha a versão:")
    print("1. CV Fiel ao PDF Original (formato idêntico)")
    print("2. CV Atualizado (com experiência real Java 6/7/8)")
    print("3. Ambas versões")
    
    try:
        escolha = input("\nDigite 1, 2 ou 3: ").strip()
        
        if escolha == "1":
            criar_cv_java_word_fiel()
        elif escolha == "2":
            criar_cv_java_atualizado()
        elif escolha == "3":
            criar_cv_java_word_fiel()
            criar_cv_java_atualizado()
            print("\n" + "="*70)
            print("✅ AMBAS VERSÕES GERADAS COM SUCESSO!")
        else:
            print("\n⚠️  Opção inválida. Gerando versão fiel ao PDF...")
            criar_cv_java_word_fiel()
    
    except KeyboardInterrupt:
        print("\n\n❌ Operação cancelada pelo usuário.")
    except Exception as e:
        print(f"\n❌ Erro: {e}")
        print("\n💡 Dica: Certifique-se de ter o python-docx instalado:")
        print("   pip install python-docx")